from docx import Document
import sys

def read_docx(path):
    try:
        doc = Document(path)
        print(f"Reading: {path}\n")
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                print(para.text)
        
        # Also check tables
        print("\n--- TABLES ---\n")
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                print(" | ".join(row_text))
            print("-" * 20)

    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    read_docx("選址評估報告 -台北市中山區松江路111號 (1).docx")
